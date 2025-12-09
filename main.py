import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import threading
import os
import subprocess

def extract_text_with_pypdf2(pdf_path):
    text_content = []
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            text_content.append(f"Page {page_num + 1}:\n{text}\n" if text else f"Page {page_num + 1}: No text found.\n")
    return text_content

def extract_images_with_pymupdf(pdf_path):
    images_content = []
    pdf_document = fitz.open(pdf_path)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        page_images = []
        
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]  # The image reference number in the PDF

            # Extract image bytes from the PDF
            base_image = pdf_document.extract_image(xref)
            img_bytes = base_image["image"]

            # Use Pillow to convert the image to a PNG format in a BytesIO stream
            img_stream = io.BytesIO(img_bytes)
            pil_image = Image.open(img_stream).convert("RGB")  # Ensure it's in RGB mode
            png_stream = io.BytesIO()
            pil_image.save(png_stream, format="PNG")  # Convert to PNG
            png_stream.seek(0)

            page_images.append(png_stream)  # Store the PNG stream
        images_content.append(page_images)

    pdf_document.close()
    return images_content

def pdf_to_word(pdf_path, word_path):
    # Create a new Word document
    doc = Document()
    doc.add_heading("PDF to Word Conversion", level=1)

    # Extract text using PyPDF2
    text_content = extract_text_with_pypdf2(pdf_path)
    
    # Extract images using PyMuPDF
    images_content = extract_images_with_pymupdf(pdf_path)

    # Add extracted text and images to the Word document
    for page_num, page_text in enumerate(text_content):
        doc.add_heading(f"Page {page_num + 1}", level=2)
        doc.add_paragraph(page_text)

        # Add images if present
        if page_num < len(images_content) and images_content[page_num]:
            for img_index, img_stream in enumerate(images_content[page_num]):
                img_paragraph = doc.add_paragraph()
                img_paragraph.add_run(f"Image {img_index + 1} on Page {page_num + 1}").bold = True
                doc.add_picture(img_stream, width=Inches(5))  # Adjust size as needed
                doc.add_paragraph("")

        # Add a page break after each page
        doc.add_page_break()

    # Save the Word document
    doc.save(word_path)
    print(f"Conversion complete! The Word document is saved at: {word_path}")

class PDFToWordConverter:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("PDF to Word Converter")
        self.window.geometry("600x400")
        self.window.configure(bg="#f0f0f0")
        
        # Variables
        self.pdf_path = tk.StringVar()
        self.word_path = tk.StringVar()
        
        self.create_widgets()
        
    def create_widgets(self):
        # Main frame
        main_frame = ttk.Frame(self.window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_label = ttk.Label(
            main_frame, 
            text="PDF to Word Converter",
            font=("Helvetica", 16, "bold")
        )
        title_label.pack(pady=20)
        
        # Create a frame for the download buttons
        self.download_frame = ttk.Frame(main_frame)
        self.download_frame.pack(fill=tk.X, pady=5)
        
        # Create download buttons (initially hidden)
        self.open_file_btn = ttk.Button(
            self.download_frame,
            text="Open Document",
            command=self.open_word_file,
            style="Accent.TButton"
        )
        
        self.open_folder_btn = ttk.Button(
            self.download_frame,
            text="Open Containing Folder",
            command=self.open_containing_folder
        )
        
        # PDF file selection
        pdf_frame = ttk.Frame(main_frame)
        pdf_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(pdf_frame, text="PDF File:").pack(side=tk.LEFT)
        ttk.Entry(pdf_frame, textvariable=self.pdf_path, width=50).pack(side=tk.LEFT, padx=5)
        ttk.Button(pdf_frame, text="Browse", command=self.browse_pdf).pack(side=tk.LEFT)
        
        # Word file selection
        word_frame = ttk.Frame(main_frame)
        word_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(word_frame, text="Save As:").pack(side=tk.LEFT)
        ttk.Entry(word_frame, textvariable=self.word_path, width=50).pack(side=tk.LEFT, padx=5)
        ttk.Button(word_frame, text="Browse", command=self.browse_word).pack(side=tk.LEFT)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=20)
        
        # Convert button
        self.convert_btn = ttk.Button(
            main_frame,
            text="Convert PDF to Word",
            command=self.start_conversion,
            style="Accent.TButton"
        )
        self.convert_btn.pack(pady=10)
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="")
        self.status_label.pack(pady=10)
        
    def browse_pdf(self):
        filename = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if filename:
            self.pdf_path.set(filename)
            # Automatically set a default word path
            base_name = os.path.splitext(filename)[0]
            self.word_path.set(f"{base_name}.docx")
    
    def browse_word(self):
        filename = filedialog.asksaveasfilename(
            title="Save Word File",
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.word_path.set(filename)
    
    def start_conversion(self):
        if not self.pdf_path.get():
            messagebox.showerror("Error", "Please select a PDF file")
            return
        if not self.word_path.get():
            messagebox.showerror("Error", "Please select where to save the Word file")
            return
            
        self.convert_btn.configure(state='disabled')
        self.progress.start()
        self.status_label.configure(text="Converting... Please wait.")
        
        # Run conversion in a separate thread
        thread = threading.Thread(target=self.convert_file)
        thread.daemon = True
        thread.start()
    
    def convert_file(self):
        try:
            pdf_to_word(self.pdf_path.get(), self.word_path.get())
            self.window.after(0, self.conversion_complete, True)
        except Exception as e:
            self.window.after(0, self.conversion_complete, False, str(e))
    
    def conversion_complete(self, success, error_message=None):
        self.progress.stop()
        self.convert_btn.configure(state='normal')
        
        if success:
            self.status_label.configure(text="Conversion completed successfully!")
            messagebox.showinfo("Success", "PDF has been converted to Word successfully!")
            # Show download buttons
            self.open_file_btn.pack(side=tk.LEFT, padx=5)
            self.open_folder_btn.pack(side=tk.LEFT, padx=5)
        else:
            self.status_label.configure(text="Conversion failed!")
            messagebox.showerror("Error", f"Failed to convert PDF: {error_message}")
            # Hide download buttons
            self.open_file_btn.pack_forget()
            self.open_folder_btn.pack_forget()
    
    def open_word_file(self):
        """Open the converted Word document"""
        if os.path.exists(self.word_path.get()):
            try:
                os.startfile(self.word_path.get())
            except AttributeError:  # For non-Windows systems
                subprocess.run(['start', self.word_path.get()], shell=True)
        else:
            messagebox.showerror("Error", "The converted file cannot be found!")
    
    def open_containing_folder(self):
        """Open the folder containing the converted Word document"""
        if os.path.exists(self.word_path.get()):
            folder_path = os.path.dirname(os.path.abspath(self.word_path.get()))
            try:
                os.startfile(folder_path)
            except AttributeError:  # For non-Windows systems
                subprocess.run(['start', folder_path], shell=True)
        else:
            messagebox.showerror("Error", "The folder cannot be found!")
    
    def run(self):
        self.window.mainloop()

if __name__ == "__main__":
    app = PDFToWordConverter()
    app.run()
